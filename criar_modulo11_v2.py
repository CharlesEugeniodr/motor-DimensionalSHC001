from __future__ import annotations
from pathlib import Path
from copy import deepcopy
import json, csv, math, os, zipfile, hashlib
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/mnt/data/M11_v2_work')
DATA=ROOT/'data'; FIG=ROOT/'figures'; CODE=ROOT/'code'
TEMPLATE=Path('/mnt/data/SHC001_Modulo_11_E3A_Validacao_Numerica_D1_P0_PT_v1.docx')
OUT=Path('/mnt/data/SHC001_Modulo_11_D1_P2_R0_Projeto_Aerodinamico_Tecnico_PT_v2.docx')

p=json.load(open(DATA/'baseline_D1_P2_R0.json',encoding='utf-8'))
nom=json.load(open(DATA/'desempenho_nominal_D1_P2_R0.json',encoding='utf-8'))
rad=pd.read_csv(DATA/'projeto_radial_D1_P2_R0.csv')
checks=pd.read_csv(DATA/'verificacoes_automatizadas.csv')
mit=pd.read_csv(DATA/'matriz_mitigacoes_BF.csv')
modal=pd.read_csv(DATA/'margem_modal_reduzida.csv')
tests=pd.read_csv(DATA/'plano_T13_T18_razao_torque.csv')

BLUE='17365D'; BLUE2='2F5597'; LIGHT='D9EAF7'; LIGHT2='EAF2F8'; GRAY='E7E6E6'; DARK='1F1F1F'; RED='C00000'; GREEN='548235'; YELLOW='FFF2CC'

def clear_body(doc):
    body=doc._element.body
    sectPr=body.sectPr
    for child in list(body):
        if child is not sectPr:
            body.remove(child)

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr()
    shd=tcPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_text_color(cell, color='FFFFFF'):
    for p0 in cell.paragraphs:
        for r in p0.runs:
            r.font.color.rgb=RGBColor.from_string(color)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for cell,w in zip(row.cells,widths_cm): cell.width=Cm(w)

def add_table(doc, headers, rows, widths=None, font_size=8.1, first_col_bold=False):
    table=doc.add_table(rows=1, cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    hdr=table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        c=hdr.cells[i]; c.text=str(h); set_cell_shading(c,BLUE); set_cell_text_color(c)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for pp in c.paragraphs:
            pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in pp.runs: r.bold=True; r.font.size=Pt(font_size); r.font.name='Arial'
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri%2: set_cell_shading(cells[i],'F7F9FB')
            for pp in cells[i].paragraphs:
                for r in pp.runs: r.font.size=Pt(font_size); r.font.name='Arial'
                if i==0 and first_col_bold:
                    for r in pp.runs: r.bold=True
            set_cell_margins(cells[i])
    if widths: set_col_widths(table,widths)
    doc.add_paragraph('')
    return table

def add_p(doc,text='',bold_prefix=None,style='Normal',align=None,space_after=4):
    para=doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r=para.add_run(bold_prefix); r.bold=True; para.add_run(text[len(bold_prefix):])
    else: para.add_run(text)
    if align is not None: para.alignment=align
    para.paragraph_format.space_after=Pt(space_after)
    para.paragraph_format.line_spacing=1.08
    for r in para.runs: r.font.name='Arial'; r.font.size=Pt(9.3)
    return para

def add_bullets(doc,items):
    for item in items:
        para=doc.add_paragraph(style='List Bullet'); para.paragraph_format.space_after=Pt(2)
        r=para.add_run(item); r.font.name='Arial'; r.font.size=Pt(9.2)

def add_numbered(doc,items):
    for item in items:
        para=doc.add_paragraph(style='List Number'); para.paragraph_format.space_after=Pt(2)
        r=para.add_run(item); r.font.name='Arial'; r.font.size=Pt(9.2)

def add_eq(doc,text):
    para=doc.add_paragraph(style='Equation'); para.alignment=WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.space_before=Pt(4); para.paragraph_format.space_after=Pt(5)
    r=para.add_run(text); r.font.name='FreeSerif'; r.font.size=Pt(11)
    return para

def add_callout(doc,title,text,fill=LIGHT2):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; c=t.cell(0,0); set_cell_shading(c,fill); set_cell_margins(c,120,140,120,140)
    pp=c.paragraphs[0]; rr=pp.add_run(title); rr.bold=True; rr.font.name='Arial'; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor.from_string(BLUE)
    rr=pp.add_run('\n'+text); rr.font.name='Arial'; rr.font.size=Pt(9.1)
    doc.add_paragraph('')

def add_picture(doc,path,caption,width=6.7):
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER; para.paragraph_format.keep_with_next=True
    para.add_run().add_picture(str(path),width=Inches(width))
    cap=doc.add_paragraph(style='Caption'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after=Pt(5); cap.paragraph_format.keep_with_next=False
    r=cap.add_run(caption); r.font.name='Arial'; r.font.size=Pt(8.4); r.italic=True

def add_heading(doc,text,level=1):
    p0=doc.add_heading(text,level=level)
    p0.paragraph_format.keep_with_next=True
    for r in p0.runs: r.font.name='Arial'
    return p0

def add_page_break(doc): doc.add_page_break()

def add_page_field(paragraph):
    run=paragraph.add_run(); fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin'); instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '; fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end'); run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def fmt(x,n=3): return f'{x:.{n}f}'.replace('.',',')

doc=Document(TEMPLATE); clear_body(doc)
sec=doc.sections[0]
sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.4); sec.left_margin=Cm(1.7); sec.right_margin=Cm(1.5)
# styles
normal=doc.styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(9.3)
for sty in ['Heading 1','Heading 2','Heading 3']:
    s=doc.styles[sty]; s.font.name='Arial'; s.font.color.rgb=RGBColor.from_string(BLUE)
# header/footer
hdr=sec.header.paragraphs[0]; hdr.clear(); hdr.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=hdr.add_run('SHC-001 | Módulo 11 — edição técnica v2 | D1-P2-R0'); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(BLUE)
ftr=sec.footer.paragraphs[0]; ftr.clear(); ftr.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ftr.add_run('Versão 2.0 — julho de 2026 | Projeto aerodinâmico preliminar e medidas mitigatórias | Página '); r.font.name='Arial'; r.font.size=Pt(8); add_page_field(ftr)

# COVER
p0=doc.add_paragraph(); p0.alignment=WD_ALIGN_PARAGRAPH.CENTER; p0.paragraph_format.space_before=Pt(28)
r=p0.add_run('SHC-001'); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(BLUE); r.font.name='Arial'
p0=doc.add_paragraph(); p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p0.add_run('MÓDULO 11 — EDIÇÃO TÉCNICA REVISADA'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor.from_string(BLUE2); r.font.name='Arial'
p0=doc.add_paragraph(); p0.alignment=WD_ALIGN_PARAGRAPH.CENTER; p0.paragraph_format.space_before=Pt(8)
r=p0.add_run('D1-P2-R0'); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor.from_string(DARK); r.font.name='Arial'
p0=doc.add_paragraph(); p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p0.add_run('PROJETO AERODINÂMICO PRELIMINAR DOS ROTORES CONTRARROTATIVOS'); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(BLUE); r.font.name='Arial'
p0=doc.add_paragraph(); p0.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p0.add_run('Mean-line, triângulos de velocidade, geometria radial das pás, torque, dinâmica e medidas mitigatórias para CFD/FEA'); r.font.size=Pt(10.5); r.font.name='Arial'
add_picture(doc,FIG/'fig03_triangulos_velocidade.png','Síntese cinemática do estágio contrarrotativo no raio médio.',width=6.5)
add_callout(doc,'CLASSIFICAÇÃO DA EVIDÊNCIA','Cálculos analíticos, projeto médio e geometria paramétrica executados. O documento não declara CFD 3D, FEA 3D, prontidão para usinagem ou liberação para ensaio pressurizado.',YELLOW)
add_p(doc,'Documento de pesquisa independente. A expressão “revisão adversarial” descreve uma metodologia de crítica técnica e não representa parecer ou endosso institucional do ITA, NASA, INPE ou qualquer órgão externo.',align=WD_ALIGN_PARAGRAPH.CENTER,space_after=0)
add_page_break(doc)

# CONTROL
add_heading(doc,'Controle documental',1)
add_table(doc,['Campo','Registro'],[
['Documento','SHC001 — Módulo 11 — D1-P2-R0 — Projeto Aerodinâmico Técnico'],
['Revisão','2.0 — edição técnica substitutiva'],
['Configuração','D1-P2-R0'],
['Natureza','Projeto preliminar de turbomáquina axial contrarrotativa de fluxo frio'],
['Fluido de referência','Ar seco a 293,15 K; densidade de projeto 1,20 kg/m³'],
['Escopo executado','Área ativa, mean-line, triângulos de velocidade, solidez, corda, torção, Re/Mach, torque/potência, envelope modal reduzido e requisitos de mitigação'],
['Fora do escopo','CFD 3D, FEA 3D, CAD nativo liberado, qualificação de materiais, fabricação, ensaio físico e certificação'],
['Relação com a v1','A v1 permanece como registro histórico da auditoria D1-P0. Esta revisão substitui a definição aerodinâmica de pás e o envelope operacional preliminar.'],
],widths=[4.0,12.0],font_size=8.4,first_col_bold=True)
add_callout(doc,'REGRA DE LEITURA','“Mitigado analiticamente” significa que existe uma formulação e uma baseline numérica verificável. Não significa que o fenômeno esteja validado por CFD ou ensaio.')

add_heading(doc,'Sumário executivo',1)
add_p(doc,'A auditoria anterior mostrou que o anel D1-P1, com área ativa de aproximadamente 0,208 m², vazão de 0,40 kg/s e rotação de 1.200 rpm, conduzia a um coeficiente de fluxo muito baixo. A presente revisão corrige a causa, não apenas a aparência da pá: introduz um cartucho anular com raios de 140 e 200 mm, reduz a rotação nominal para 600 rpm e preserva 0,40 kg/s. O ponto nominal passa a apresentar φ = 0,480, razão cubo–ponta de 0,70 e velocidades relativas de baixo Mach.')
add_p(doc,'Com a cinemática congelada, foram derivados os ângulos de entrada e saída de R1 e R2 em três estações radiais. Foram fixadas 12 pás por rotor, cordas de 80 a 100 mm e solidez de 0,955 a 1,091. A geometria deixa de ser uma superfície helicoidal arbitrária e passa a ser definida por leis radiais de ângulo metálico, perfil de espessura relativa de 10%, bordos fabricáveis e folga nominal de 0,50 mm.')
add_p(doc,'A potência ideal transferida ao fluido no ponto nominal é aproximadamente 51,8 W, com torque ideal de 0,412 N·m por rotor. Esses valores servem para dimensionar instrumentação e motores, mas ainda não incluem perdas, atrito, inércia ou margem de controle. A condição de torque estrutural mínimo será identificada em bancada por varredura de |N₂|/N₁ entre 0,80 e 1,00; o fechamento reduzido anterior indica um ponto próximo de 0,88.')
add_p(doc,'O limite operacional permanece em 1.200 rpm. A sobrevelocidade de ensaio é proposta em 1.400 rpm somente após FEA modal tridimensional e revisão de prontidão. O valor de 1.500 rpm permanece como caso de projeto para FEA, não como autorização operacional.')
add_page_break(doc)

# CONTENTS
add_heading(doc,'Índice técnico',1)
contents=[
'1. Antecedentes e problema central','2. Matriz de medidas mitigatórias','3. Baseline D1-P2-R0','4. Redimensionamento do anel ativo','5. Metodologia mean-line','6. Triângulos de velocidade','7. Número de pás, corda, espaçamento e solidez','8. Ângulos metálicos e torção radial','9. Perfil, espessura, curvatura e bordos','10. Incidência, desvio e critérios aerodinâmicos','11. Folgas, Reynolds e compressibilidade','12. Torque, potência e cancelamento de reação','13. Estator e sistema de saída','14. Envelope modal, BPF e limites de rotação','15. Pressurização e proteção contra sobrepressão','16. Motores, eixos, mancais e balanceamento','17. Orçamento de massa e interfaces CAD','18. Requisitos de CFD 3D','19. Requisitos de FEA e rotodinâmica','20. Plano de ensaio atualizado','21. Critérios de aceitação e gates','22. Resultados das verificações automatizadas','23. Conclusões e estado técnico','Apêndices A–E.'
]
add_numbered(doc,contents)
add_page_break(doc)

# 1
add_heading(doc,'1. Antecedentes e problema central',1)
add_p(doc,'O D1-P0 e o D1-P1 foram importantes para auditar volumes, bypass, cascas, eixos e interfaces. Entretanto, suas pás foram geradas como superfícies helicoidais de espessura quase constante. Essa geometria não possuía uma definição completa de corda, solidez, incidência, desvio, perfil ou carregamento de estágio. Assim, ela era adequada para ocupação espacial e discussão de arquitetura, mas insuficiente para análise aerodinâmica de turbomáquina.')
add_p(doc,'A inconsistência dominante era a relação entre área de passagem, vazão e rotação. Para um rotor axial, a grandeza de primeira ordem é o coeficiente de fluxo:')
add_eq(doc,'φ = Vₓ / U')
add_p(doc,'em que Vₓ é a velocidade axial média e U = ωr é a velocidade periférica. Um φ muito baixo indica que o rotor gira rapidamente diante de um escoamento axial lento; os ângulos relativos tornam-se extremos e a incidência fica altamente sensível a pequenas variações de vazão.')
add_picture(doc,FIG/'fig01_aneis_D1P1_D1P2.png','Figura 1 — Comparação entre o anel D1-P1 e o cartucho ativo D1-P2-R0.')
add_callout(doc,'DECISÃO DE PROJETO','Manter o envelope externo da bancada e inserir um cartucho intercambiável com área ativa menor. Isso permite estudar diferentes estágios sem reconstruir toda a estrutura.')

# 2
add_heading(doc,'2. Matriz de medidas mitigatórias',1)
rows=[]
for _,r in mit.iterrows(): rows.append([r.codigo,r.problema,r.estado,r.mitigacao_atual,r.criterio_fechamento])
add_table(doc,['Código','Problema','Estado','Mitigação nesta revisão','Fechamento exigido'],rows,widths=[1.5,3.4,2.3,4.4,4.7],font_size=7.3,first_col_bold=True)
add_picture(doc,FIG/'fig10_mitigacoes.png','Figura 2 — Estado dos bloqueios de fabricação após o projeto D1-P2-R0.')
add_p(doc,'Apenas BF-02 recebeu mitigação analítica direta nesta revisão. Os demais bloqueios receberam requisitos e interfaces, mas continuam abertos até a produção da evidência correspondente. Em especial, BF-03 e BF-05 permanecem bloqueadores de rotação e pressurização.')

# 3
add_heading(doc,'3. Baseline D1-P2-R0',1)
base_rows=[
['Fluido','Ar seco'],['Temperatura de referência','293,15 K'],['Densidade de projeto','1,20 kg/m³'],['Viscosidade dinâmica','1,81×10⁻⁵ Pa·s'],['Vazão nominal','0,40 kg/s'],['Raio do cubo','140 mm'],['Raio da ponta','200 mm'],['Altura anular','60 mm'],['Área ativa',f"{fmt(p['area_m2'],5)} m²"],['Razão cubo–ponta',fmt(p['hub_tip_ratio'],2)],['Rotação nominal','600 rpm'],['Rotação operacional máxima','1.200 rpm'],['Sobrevelocidade de ensaio proposta','1.400 rpm, condicionada à FEA'],['Caso de projeto FEA','1.500 rpm'],['Número de pás','12 em R1 e 12 em R2'],['Folga nominal de ponta','0,50 mm'],['Razão de swirl de projeto','Vθ/U = 0,55'],['Desvio inicial de estudo','|δ| = 4°'],]
add_table(doc,['Parâmetro','Valor'],base_rows,widths=[7.0,9.0],font_size=8.4,first_col_bold=True)
add_callout(doc,'HIPÓTESE DE DENSIDADE','A densidade de 1,20 kg/m³ é utilizada para fechar a primeira baseline. Quando a pressão estática e a temperatura locais do rotor forem definidas, Vₓ, φ, Reynolds e os ângulos deverão ser recalculados.')

# 4
add_heading(doc,'4. Redimensionamento do anel ativo',1)
add_eq(doc,'A = π(rₜ² − rₕ²) = 0,064088 m²')
add_eq(doc,'Vₓ = ṁ/(ρA) = 5,201 m/s')
add_eq(doc,'rₘ = √[(rₕ² + rₜ²)/2] = 0,172627 m')
add_eq(doc,'Uₘ = 2πN rₘ/60 = 10,846 m/s')
add_eq(doc,'φ = Vₓ/Uₘ = 0,4795')
add_p(doc,'O valor de φ não deve ser interpretado como “ótimo universal”. Ele é um ponto de partida tratável para uma máquina axial de baixo carregamento. A faixa de operação deverá ser mapeada por CFD e ensaio, porque o limite de incidência depende do perfil, da solidez, do Reynolds e das folgas.')
add_picture(doc,FIG/'fig02_mapa_phi.png','Figura 3 — Mapa paramétrico de φ em função da vazão e da rotação.')

# 5
add_heading(doc,'5. Metodologia mean-line',1)
add_p(doc,'O projeto médio é executado em três estações — cubo, raio médio quadrático e ponta — e segue a sequência:')
add_numbered(doc,[
'Fixar o canal anular, a vazão e a rotação.',
'Calcular Vₓ e U(r).',
'Definir uma lei inicial para Vθ na saída de R1.',
'Construir os triângulos absolutos e relativos de R1.',
'Construir os triângulos de R2 com velocidade periférica de sinal oposto.',
'Converter os ângulos de escoamento em ângulos metálicos por uma hipótese explícita de incidência e desvio.',
'Dimensionar corda e número de pás para obter solidez radial controlada.',
'Gerar uma superfície tridimensional por seções radiais e linha média contínua.',
'Validar carregamento, desvio, perdas e separação por CFD 3D.',
])
add_eq(doc,'V⃗ = U⃗ + W⃗')
add_p(doc,'O projeto utiliza sinais explícitos para o componente tangencial. R1 gira no sentido positivo e R2 no sentido negativo. Os ângulos são medidos a partir do eixo axial; valores negativos e positivos indicam lados opostos do plano meridional.')

# 6
add_heading(doc,'6. Triângulos de velocidade',1)
add_p(doc,'A lei inicial de R1 é Vθ,2 = 0,55U. Na saída de R2, o objetivo mean-line é Vθ,3 ≈ 0 no raio médio. Essa escolha cria um estágio de baixo carregamento adequado para uma bancada e não pretende maximizar razão de pressão.')
add_eq(doc,'β = arctan[(Vθ − U)/Vₓ]')
add_picture(doc,FIG/'fig03_triangulos_velocidade.png','Figura 4 — Triângulos de velocidade de R1 e R2 no raio médio.')
flow_rows=[]
for _,r in rad.iterrows():
    flow_rows.append([r.estacao,fmt(r.r_m*1000,1),fmt(r.U_m_s,2),fmt(r.Vx_m_s,2),fmt(r.Vtheta_R1_saida_m_s,2),fmt(r.beta_R1_entrada_deg,1),fmt(r.beta_R1_saida_deg,1),fmt(r.beta_R2_entrada_deg,1),fmt(r.beta_R2_saida_deg,1)])
add_table(doc,['Estação','r (mm)','U','Vₓ','Vθ R1 saída','β R1 ent.','β R1 saí.','β R2 ent.','β R2 saí.'],flow_rows,widths=[1.7,1.4,1.3,1.3,1.8,1.6,1.6,1.6,1.6],font_size=7.2)
add_callout(doc,'LIMITAÇÃO','O fechamento Vθ,3 ≈ 0 é aplicado ao projeto médio. Camada limite, escoamentos secundários, folga de ponta e interação rotor–rotor gerarão swirl residual tridimensional.')

# 7
add_heading(doc,'7. Número de pás, corda, espaçamento e solidez',1)
add_p(doc,'A versão anterior continha seis elementos por rotor como placeholders geométricos. Esta revisão fixa 12 pás por rotor para obter solidez próxima de 1 com cordas fabricáveis. O espaçamento circunferencial e a solidez são:')
add_eq(doc,'s(r) = 2πr/Z     ;     σ(r) = c(r)/s(r)')
sol_rows=[]
for _,r in rad.iterrows(): sol_rows.append([r.estacao,fmt(r.r_m*1000,1),fmt(r.espacamento_m*1000,1),fmt(r.corda_m*1000,1),fmt(r.solidez,3)])
add_table(doc,['Estação','Raio (mm)','Espaçamento (mm)','Corda (mm)','Solidez'],sol_rows,widths=[3.0,2.5,3.2,2.5,2.3],font_size=8.2)
add_picture(doc,FIG/'fig04_corda_solidez.png','Figura 5 — Lei radial de corda e solidez.')
add_p(doc,'O número de pás não está “validado” pelo cálculo de solidez. Ele também afeta bloqueio, frequência de passagem, interação potencial, fabricação, massa e estabilidade. Z = 12 é uma baseline a ser comparada, no mínimo, com Z = 10 e Z = 14 em estudos posteriores.')

# 8
add_heading(doc,'8. Ângulos metálicos e torção radial',1)
add_p(doc,'A incidência nominal é definida inicialmente como zero. O desvio de 4° é um parâmetro de partida, não uma correlação validada. Os ângulos metálicos resultantes são:')
angle_rows=[]
for _,r in rad.iterrows(): angle_rows.append([r.estacao,fmt(r.metal_R1_entrada_deg,1),fmt(r.metal_R1_saida_deg,1),fmt(r.cambra_R1_deg,1),fmt(r.metal_R2_entrada_deg,1),fmt(r.metal_R2_saida_deg,1),fmt(r.cambra_R2_deg,1)])
add_table(doc,['Estação','R1 metal ent.','R1 metal saí.','Cambra R1','R2 metal ent.','R2 metal saí.','Cambra R2'],angle_rows,widths=[2.3,2.2,2.2,2.0,2.2,2.2,2.0],font_size=7.8)
add_picture(doc,FIG/'fig05_angulos_metalicos.png','Figura 6 — Leis radiais preliminares dos ângulos metálicos.')
add_p(doc,'A torção tridimensional deve ser produzida por seções em cubo, meio e ponta, interpoladas por função suave. Não existe mais um único “ângulo helicoidal de 62°”. Cada bordo possui uma lei radial própria.')

# 9
add_heading(doc,'9. Perfil, espessura, curvatura e bordos',1)
add_p(doc,'A família NACA 65 de compressor é adotada apenas como referência de cascata e ponto de partida para seleção de cambra e incidência. A designação “NACA 65 com 10%” não deve ser usada como geometria final sem a geração das ordenadas canônicas ou de uma seção paramétrica equivalente.')
add_p(doc,'Para o primeiro CAD analisável, recomenda-se uma seção paramétrica com linha média cúbica, tangentes iguais aos ângulos metálicos e distribuição de espessura de 10% da corda. Essa seção deverá receber um identificador próprio, evitando atribuição indevida a uma seção NACA específica.')
profile_rows=[]
for _,r in rad.iterrows(): profile_rows.append([r.estacao,fmt(r.corda_m*1000,1),fmt(r.tmax_m*1000,1),f"{fmt(r.r_LE_min_m*1000,1)}–{fmt(r.r_LE_max_m*1000,1)}",'0,8–1,2'])
add_table(doc,['Estação','Corda (mm)','t máximo (mm)','Raio LE (mm)','Espessura TE (mm)'],profile_rows,widths=[3.0,3.0,3.0,3.5,3.5],font_size=8.2)
add_bullets(doc,[
'Raiz com filete inicial de 8 a 12 mm, sujeito à FEA.',
'Bordo de ataque arredondado, sem aresta de usinagem.',
'Bordo de fuga truncado e inspecionável.',
'Empilhamento preferencial próximo ao quarto de corda ou centroide estrutural, a ser comparado em FEA.',
'Transição contínua de curvatura entre estações, sem quinas na superfície loft.',
])

# 10
add_heading(doc,'10. Incidência, desvio e critérios aerodinâmicos',1)
add_eq(doc,'i = βfluxo,entrada − βmetal,entrada')
add_eq(doc,'δ = βfluxo,saída − βmetal,saída')
add_p(doc,'A campanha CFD deve varrer incidência entre −6° e +6° e desvio efetivo entre 2° e 8°. O ponto nominal de i = 0° e |δ| = 4° é utilizado apenas para construir a primeira geometria. As correlações de cascata NACA 65 podem orientar a seleção, mas a baixa faixa de Reynolds e a dupla rotação exigem validação específica.')
add_p(doc,'Os critérios de projeto aerodinâmico deverão incluir: fator de difusão, razão de velocidades relativas, perda de pressão total, coeficiente de desvio, carregamento superficial, margem de separação e sensibilidade à folga. Nenhum desses critérios foi ainda fechado tridimensionalmente.')

# 11
add_heading(doc,'11. Folgas, Reynolds e compressibilidade',1)
add_eq(doc,'g/h = 0,50/60 = 0,833%')
add_p(doc,'A folga de 0,50 mm satisfaz a meta geométrica inferior a 1% do span, mas sua viabilidade depende da cadeia de tolerâncias, do batimento, da deformação centrífuga, da flexão das pás e da expansão térmica. O valor final não deve ser congelado sem FEA e análise de montagem.')
add_picture(doc,FIG/'fig06_re_mach.png','Figura 7 — Reynolds e Mach relativo no envelope paramétrico.')
re_rows=[]
for _,r in rad.iterrows(): re_rows.append([r.estacao,fmt(r.Wmax_m_s,2),f"{r.Re_corda_nom:.2e}",fmt(r.Mach_rel_nom,4)])
add_table(doc,['Estação','W máximo a 600 rpm (m/s)','Re da corda','Mach relativo'],re_rows,widths=[3.5,4.0,3.5,3.5],font_size=8.2)
add_p(doc,'No ponto nominal, os Reynolds locais ficam aproximadamente entre 6×10⁴ e 1,2×10⁵. No limite de 1.200 rpm, o valor de ponta se aproxima de 2,4×10⁵ e o Mach relativo máximo calculado permanece em torno de 0,115. O tratamento de baixo Mach é adequado, mas o regime de transição viscosa deve ser modelado ou analisado por sensibilidade.')

# 12
add_heading(doc,'12. Torque, potência e cancelamento de reação',1)
add_eq(doc,'ΔhR1 = UΔVθ = 64,71 J/kg')
add_eq(doc,'τ = ṁ rₘ ΔVθ = 0,412 N·m por rotor')
add_eq(doc,'P = τω = 25,88 W por rotor')
add_eq(doc,'Pfluido,total ≈ 51,76 W')
add_p(doc,'Os valores são ideais. O motor deve suprir também perdas nos mancais, vedações, acoplamentos, arrasto de disco, aceleração da inércia e reserva de controle. A potência nominal do motor não deve ser selecionada multiplicando apenas 26 W por um fator arbitrário; é necessário um orçamento de perdas e um perfil de aceleração.')
add_p(doc,'A igualdade de rpm não garante reação líquida nula. O Módulo 11 v1 mostrou, em um fechamento reduzido com efetividade ε = 0,12, uma razão aproximada |N₂|/N₁ = 1 − ε = 0,88. Essa relação não é uma constante física do SHC; ela é um centro de varredura experimental.')
add_picture(doc,FIG/'fig08_razao_torque.png','Figura 8 — Varredura proposta da razão de rotação para minimizar torque estrutural.')
add_callout(doc,'REQUISITO DE CONTROLE','A malha de contrarrotação deverá usar torque medido em R1 e R2. O objetivo não é manter rpm iguais, mas minimizar o torque transmitido à estrutura dentro do envelope de operação.')

# 13
add_heading(doc,'13. Estator e sistema de saída',1)
add_p(doc,'Mesmo quando o projeto médio de R2 impõe Vθ ≈ 0, o escoamento real terá swirl residual por não uniformidade radial, camada limite, folga e interação entre fileiras. O estator deve ser intercambiável e cumprir três funções:')
add_bullets(doc,[
'Homogeneizar a direção e o módulo da velocidade na saída.',
'Recuperar pressão estática a partir de swirl residual e difusão controlada.',
'Promover transição suave para o difusor ou elemento de descarga.',
])
add_p(doc,'A carga inicial prevista para o estator é modesta, com ΔVθ/U entre 0,15 e 0,20 no envelope de projeto. O valor final será obtido após CFD de R1+R2, evitando superdimensionar o estator e introduzir perdas desnecessárias.')
add_picture(doc,FIG/'fig09_estator_envelope.png','Figura 9 — Envelope inicial de carga previsto para o estator.')
add_p(doc,'O bocal anular anterior de aproximadamente 0,163 m² não deve ser transferido automaticamente ao cartucho de 0,064 m². Para fluxo frio, a saída deve ser modular: duto de área constante, difusor ou restrição ajustável. A comparação A/B/C deve preservar a mesma condição de saída ou registrar explicitamente a diferença.')

# 14
add_heading(doc,'14. Envelope modal, BPF e limites de rotação',1)
add_eq(doc,'fBPF = ZN/60')
add_picture(doc,FIG/'fig07_bpf_modal.png','Figura 10 — Frequência de passagem das pás e modo reduzido de referência.')
modal_rows=[]
for _,r in modal.iterrows(): modal_rows.append([int(r.rpm),fmt(r.BPF_Hz,1),fmt(r.f_modo_reduzido_Hz,1),fmt(r.margem_pct,1)+'%'])
add_table(doc,['Rotação (rpm)','BPF 12× (Hz)','Modo reduzido (Hz)','Margem reduzida'],modal_rows,widths=[3.4,3.4,3.4,3.4],font_size=8.3)
add_p(doc,'A frequência de 337 Hz deriva de uma pá equivalente em balanço e não representa o conjunto. A rotação pode elevar algumas frequências por enrijecimento centrífugo, mas efeitos giroscópicos, flexibilidade dos cubos, mancais e base podem introduzir novos modos. Não é correto assumir que 337 Hz seja necessariamente um limite inferior do conjunto completo.')
add_table(doc,['Limite','Valor','Condição'],[
['Rotação nominal','600 rpm','Projeto mean-line'],['Máxima operacional provisória','1.200 rpm','Somente após FEA e comissionamento progressivo'],['Sobrevelocidade de ensaio proposta','1.400 rpm','Somente após Campbell, contenção e TRR'],['Caso de projeto estrutural','1.500 rpm','Carga para FEA; não é autorização de teste'],
],widths=[5.0,3.0,8.0],font_size=8.3,first_col_bold=True)

# 15
add_heading(doc,'15. Pressurização e proteção contra sobrepressão',1)
add_picture(doc,FIG/'fig11_cadeia_pressao.png','Figura 11 — Arquitetura mínima de proteção contra sobrepressão.')
add_p(doc,'A fonte, o regulador, o restritor e o plenum devem possuir pressões funcionais distintas. A proteção deve considerar o cenário de regulador travado aberto e a maior vazão que a fonte pode fornecer. A área do dispositivo de alívio não pode ser calculada apenas a partir dos 0,40 kg/s nominais.')
add_eq(doc,'ṁalívio,capacidade ≥ ṁfonte,máxima no pior caso')
add_bullets(doc,[
'Definir pressão normal, pressão máxima operacional, pressão de projeto e pressão máxima admissível.',
'Selecionar dispositivo mecânico independente da lógica de controle.',
'Prever sensor eletrônico, manômetro independente, trip e descarga para área segura.',
'Calcular contrapressão, regime crítico/subcrítico e acumulação durante o alívio.',
'Submeter carcaça, flanges, aberturas e juntas à análise estrutural correspondente.',
])
add_callout(doc,'BLOQUEIO BF-05','Nenhum ensaio pressurizado deve ser liberado antes do dimensionamento do alívio, verificação da resistência do circuito e revisão por profissional responsável pela instalação.')

# 16
add_heading(doc,'16. Motores, eixos, mancais e balanceamento',1)
add_p(doc,'O dimensionamento do acionamento deve partir de um balanço de torque e energia, não apenas do torque aerodinâmico ideal. Para cada rotor devem ser definidos:')
add_bullets(doc,[
'Torque contínuo, torque de pico, tempo de aceleração e inércia refletida.',
'Controle de velocidade e de torque, frenagem e resposta a perda de sincronismo.',
'Acoplamento, alinhamento, rigidez torsional e limite de desalinhamento.',
'Cargas radiais e axiais dos mancais, vida L10, pré-carga, temperatura e rigidez dinâmica.',
'Planos de balanceamento, desbalanceamento residual e procedimento após montagem.',
])
add_p(doc,'A ISO 21940-11 é aplicável a procedimentos e tolerâncias de rotores com comportamento rígido. Se o Campbell indicar comportamento flexível no envelope, a avaliação deverá migrar para princípios compatíveis com a ISO 21940-12. A classificação só pode ser feita após o modelo rotodinâmico.')

# 17
add_heading(doc,'17. Orçamento de massa e interfaces CAD',1)
add_p(doc,'A massa de 95,8 kg do D1-P1 é uma massa geométrica parcial. O D1-P2-R0 deve possuir orçamento de massa por grupos e margem de crescimento. O documento mestre de fabricação deverá ser CAD sólido nativo e desenhos 2D; STL permanecerá apenas como formato de visualização ou malha.')
add_table(doc,['Grupo','Método de obtenção','Status'],[
['Cartucho, liner e pás','CAD sólido com material','A recalcular'],['Cubos e eixos','CAD + material + tolerâncias','A recalcular'],['Mancais','Massa de catálogo','Aberto'],['Motores e inversores','Massa de catálogo','Aberto'],['Torquímetros e acoplamentos','Massa de catálogo','Aberto'],['Base, contenção e suportes','CAD/FEA','Aberto'],['Tubulação, válvulas e alívio','P&ID + catálogo','Aberto'],['Instrumentação e cabos','Lista de I/O + catálogo','Aberto'],['Margem de crescimento','15–30% conforme maturidade','Aberto'],
],widths=[5.0,6.0,4.5],font_size=8.1)
add_bullets(doc,[
'Entregar STEP AP242 ou equivalente, além do arquivo nativo.',
'Criar datums A/B/C e cadeia de cotas funcionais.',
'Definir ajustes de eixos, mancais e cubos.',
'Especificar concentricidade, batimento, paralelismo e posição.',
'Definir rugosidade de superfícies aerodinâmicas e assentos.',
'Gerar desenhos de inspeção e relatório de propriedades de massa.',
])

# 18
add_heading(doc,'18. Requisitos de CFD 3D',1)
add_p(doc,'O CFD MRF é o próximo nível necessário para verificar se a geometria produz a mudança de momento angular prevista. O domínio deve incluir R1, espaço entre rotores, R2, estator opcional, folgas e duto de saída.')
add_table(doc,['ID','Requisito CFD','Métrica de saída'],[
['CFD-01','Três níveis de malha e verificação de independência','Torque, Δpt, swirl e uniformidade'],['CFD-02','Modelo de baixo Mach com sensibilidade de transição/turbulência','Perdas e desvio'],['CFD-03','MRF para baseline e AMI/sliding mesh para caso selecionado','Interação temporal'],['CFD-04','Configurações A/B/C com entradas equivalentes','Benefício líquido'],['CFD-05','Folga nominal e variações de tolerância','Vazamento e perda'],['CFD-06','Varredura de incidência −6° a +6°','Margem de operação'],['CFD-07','Varredura de Z e perfil','Solidez e perdas'],['CFD-08','Balanço de massa e momento angular','Fechamento global'],
],widths=[1.7,8.4,5.4],font_size=8.1,first_col_bold=True)
add_eq(doc,'τfluido = ṁ(rVθ)saída − ṁ(rVθ)entrada')
add_p(doc,'O torque integrado nas superfícies deve ser confrontado com o balanço de momento angular. Divergências acima da incerteza numérica deverão bloquear a interpretação dos resultados.')

# 19
add_heading(doc,'19. Requisitos de FEA e rotodinâmica',1)
add_p(doc,'A FEA deve utilizar a geometria de fabricação, propriedades materiais rastreáveis e condições de apoio baseadas em mancais reais. O modelo de pá isolada permanece apenas como triagem.')
add_table(doc,['ID','Análise','Conteúdo mínimo'],[
['FEA-01','Estática centrífuga','Pás, raízes, cubos e anéis em 600/1.200/1.500 rpm'],['FEA-02','Pressão e torque','Campos CFD ou envelopes conservadores'],['FEA-03','Modal pré-tensionada','Enrijecimento centrífugo e contato'],['FEA-04','Rotodinâmica','Eixos, mancais, giroscopia e amortecimento'],['FEA-05','Campbell','1×, 2×, BPF e harmônicos relevantes'],['FEA-06','Resposta harmônica','Desbalanceamento e forças periódicas'],['FEA-07','Base e contenção','Modos estruturais e energia de fragmentos'],['FEA-08','Fadiga preliminar','Ciclos de partida, varredura e operação'],
],widths=[1.7,4.0,9.8],font_size=8.0,first_col_bold=True)
add_eq(doc,'M q¨ + (C + ΩG) q˙ + K q = F(t)')
add_p(doc,'A meta de separação modal deve ser definida no plano de projeto. O cálculo reduzido de 20% a 1.400 rpm é apenas uma triagem; não pode ser usado como certificado de margem.')

# 20
add_heading(doc,'20. Plano de ensaio atualizado',1)
add_p(doc,'A campanha de torque deve transformar o resultado do modelo reduzido em requisito de controle. A condição inicial é N1 = +600 rpm e N2 = −600 rpm, seguida da varredura T13–T18.')
trows=[]
for _,r in tests.iterrows(): trows.append([r.teste,int(r.rpm_R1),fmt(r.razao_abs_R2_R1,2),int(round(r.rpm_R2)),r.objetivo])
add_table(doc,['Teste','R1 (rpm)','|R2|/R1','R2 (rpm)','Objetivo'],trows,widths=[1.8,2.0,2.0,2.0,8.0],font_size=8.1)
add_p(doc,'Em cada ponto deverão ser medidos torque de R1, torque de R2, torque de reação da carcaça, rotação, vazão, pressão total e estática, potência elétrica e swirl residual. O ponto ótimo será definido pelo mínimo torque estrutural sujeito aos limites de perda e potência.')
add_table(doc,['Fase','Condição','Critério de avanço'],[
['T00–T04','Sem rotação; estanqueidade e instrumentação','Calibração e ausência de vazamento não previsto'],['T05–T08','Rotação sem fluxo em baixa velocidade','Vibração e temperatura dentro do limite'],['T09–T12','Fluxo com R1 isolado e R2 isolado','Torques e perdas reproduzíveis'],['T13–T18','Varredura de razão R2/R1','Identificar mínimo de torque estrutural'],['T19–T22','Comparação A/B/C','Mesmas condições de entrada e saída'],['T23–T25','Envelope progressivo','Somente após revisão dos resultados'],
],widths=[3.0,6.5,6.0],font_size=8.1)

# 21
add_heading(doc,'21. Critérios de aceitação e gates',1)
add_table(doc,['Gate','Evidência obrigatória','Resultado'],[
['G11-1 — Congelamento analítico','Baseline, equações, scripts e testes aprovados','Atendido nesta revisão'],['G11-2 — CAD analisável','STEP/nativo, seções de pá, datums e propriedades de massa','Aberto'],['G11-3 — CFD MRF','Convergência, torques, perdas e swirl','Aberto'],['G11-4 — FEA/rotodinâmica','Campbell, tensões e margens','Aberto'],['G11-5 — Segurança pneumática','P&ID, alívio e análise de pressão','Aberto'],['G11-6 — TRR','Instrumentação, contenção, procedimentos e riscos','Aberto'],['G11-7 — Bancada A/B/C','Resultados repetíveis com incerteza','Aberto'],
],widths=[4.0,8.0,3.5],font_size=8.1,first_col_bold=True)
add_callout(doc,'REGRA DE LIBERAÇÃO','A compra de matéria-prima e a usinagem de peças rotativas somente devem ocorrer após G11-2, G11-3, G11-4 e G11-5. Protótipos não rotativos de impressão 3D podem ser usados antes para montagem e visualização de interfaces.')

# 22
add_heading(doc,'22. Resultados das verificações automatizadas',1)
ckrows=[]
for _,r in checks.iterrows(): ckrows.append([r.verificacao,r.resultado,str(r.valor),r.criterio])
add_table(doc,['Verificação','Resultado','Valor','Critério'],ckrows,widths=[6.0,2.5,3.0,4.0],font_size=8.0)
add_p(doc,'Os testes verificam a coerência interna da baseline e dos cálculos implementados. Eles não validam o perfil aerodinâmico nem as margens estruturais. Dez de dez verificações foram aprovadas.')

# 23
add_heading(doc,'23. Conclusões e estado técnico',1)
add_p(doc,'O D1-P2-R0 resolve o problema de primeira ordem do coeficiente de fluxo e fornece uma definição aerodinâmica reproduzível dos rotores. O estágio passa a possuir área, vazão, rotação, triângulos de velocidade, cordas, solidez, ângulos metálicos, espessura, bordos, folga e envelope adimensional definidos.')
add_p(doc,'Isso representa uma transição real de superfície helicoidal genérica para pré-projeto de turbomáquina. Não representa, ainda, prontidão para fabricação. O desvio, as perdas, o carregamento superficial, a interação contrarrotativa e a estabilidade modal precisam ser determinados por CFD e FEA.')
add_callout(doc,'ESTADO FORMAL','D1-P2-R0: projeto mean-line e geometria aerodinâmica preliminar executados. CAD de fabricação, CFD 3D, FEA 3D, proteção pneumática e ensaio físico permanecem gates obrigatórios.')
add_picture(doc,FIG/'fig12_sequencia_liberacao.png','Figura 12 — Caminho técnico remanescente até o ensaio A/B/C.')

# APPENDICES
add_page_break(doc)
add_heading(doc,'Apêndice A — Tabela radial completa',1)
rows=[]
for _,r in rad.iterrows():
    rows.append([r.estacao,fmt(r.r_m,5),fmt(r.corda_m,3),fmt(r.espacamento_m,4),fmt(r.solidez,4),fmt(r.U_m_s,3),fmt(r.Vtheta_R1_saida_m_s,3),fmt(r.Wmax_m_s,3),f'{r.Re_corda_nom:.3e}',fmt(r.Mach_rel_nom,4)])
add_table(doc,['Estação','r (m)','c (m)','s (m)','σ','U','Vθ','Wmax','Re','Mrel'],rows,widths=[1.8,1.5,1.4,1.4,1.3,1.4,1.4,1.4,2.1,1.4],font_size=7.4)

add_heading(doc,'Apêndice B — Equações de referência',1)
for eq in [
'A = π(rₜ² − rₕ²)',
'Vₓ = ṁ/(ρA)',
'U(r) = 2πNr/60',
'φ = Vₓ/U',
's = 2πr/Z',
'σ = c/s',
'β = arctan[(Vθ − U)/Vₓ]',
'Re = ρWc/μ',
'Mrel = W/a',
'τ = ṁ Δ(rVθ)',
'P = τω',
'fBPF = ZN/60',
'Erot = ½Iω²',
]: add_eq(doc,eq)

add_heading(doc,'Apêndice C — Arquivos de reprodutibilidade',1)
add_table(doc,['Arquivo','Função'],[
['calcular_d1p2r0.py','Recalcula baseline, mapas, tabelas, figuras e verificações.'],
['baseline_D1_P2_R0.json','Parâmetros congelados desta revisão.'],
['projeto_radial_D1_P2_R0.csv','Geometria e cinemática nas três estações radiais.'],
['mapa_operacional_D1_P2_R0.csv','Vazão, rpm, φ, Re, Mach e BPF.'],
['varredura_razao_R2_R1.csv','Fechamento reduzido de torque.'],
['margem_modal_reduzida.csv','BPF e margem baseada no modo reduzido.'],
['plano_T13_T18_razao_torque.csv','Plano de varredura de contrarrotação.'],
['verificacoes_automatizadas.csv','Resultados dos dez testes internos.'],
],widths=[6.0,9.5],font_size=8.2)

add_heading(doc,'Apêndice D — Referências técnicas',1)
refs=[
'HERRIG, L. J.; EMERY, J. C.; ERWIN, J. R. Systematic Two-Dimensional Cascade Tests of NACA 65-Series Compressor Blades at Low Speeds. NACA RM L51G31 / dados consolidados em NACA TN 3916.',
'FELIX, A. R. Summary of 65-Series Compressor-Blade Low-Speed Cascade Data by Use of the Carpet-Plotting Technique. NACA, 1957.',
'LIEBLEIN, S.; SCHWENK, F. C.; BRODERICK, R. L. Diffusion Factor for Estimating Losses and Limiting Blade Loading in Axial-Flow-Compressor Blade Elements. NACA RM E53D01, 1953.',
'VERES, J. P. Axial and Centrifugal Compressor Mean Line Flow Analysis Method. NASA/TM—2009-215585, 2009.',
'ISO 21940-11:2016. Mechanical vibration — Rotor balancing — Part 11: Procedures and tolerances for rotors with rigid behaviour; Amendment 1:2022.',
'ISO 21940-12:2016. Mechanical vibration — Rotor balancing — Part 12: Procedures and tolerances for rotors with flexible behaviour.',
'ASME BPVC Section XIII, edição 2025. Rules for Overpressure Protection.',
]
for i,ref in enumerate(refs,1):
    add_p(doc,f'{i}. {ref}',space_after=2)

add_heading(doc,'Apêndice E — Registro de decisões',1)
add_table(doc,['Decisão','Justificativa','Revisão futura'],[
['Anel 140–200 mm','Elevar φ para aproximadamente 0,48','Comparar cartuchos alternativos'],['600 rpm nominal','Reduzir ângulos relativos e BPF','Mapear 300–1.200 rpm'],['12 pás','Solidez próxima de 1','Comparar 10/12/14 pás'],['Vθ/U = 0,55','Carregamento baixo e mensurável','Otimizar por CFD'],['t/c = 10%','Rigidez e seção de compressor inicial','Refinar perfil e massa'],['δ = 4°','Hipótese de construção da primeira geometria','Calibrar por cascata/CFD'],['0,50 mm de folga','Meta geométrica inicial','Fechar cadeia de tolerâncias'],['1.400 rpm overspeed','Triagem modal reduzida de 20%','Substituir por Campbell real'],
],widths=[4.0,6.5,5.0],font_size=8.1)

# Save
doc.save(OUT)
print(OUT)
